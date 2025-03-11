import os
import sys
import ollama
from pathlib import Path
from PyPDF2 import PdfReader
import markdown
import docx
import subprocess
import tempfile
import json
import glob
import numpy as np
from datetime import datetime

# 🚀 AVAILABLE COMMANDS:
COMMANDS = """
   📖 open_pdf <file>   → Process a single PDF file
   📂 open_dir <dir>    → Process all PDFs in a directory
   💬 chat              → Chat with processed PDFs
   📊 list_docs         → List all processed documents
   🔍 check_deps        → Check for required dependencies
   ❌ exit              → Quit the program
"""

# 🔥 AI MODELS
EMBEDDING_MODEL = "mxbai-embed-large"
CHAT_MODEL = "llama3.1:latest"

# 📂 File Storage
PROCESSED_DIR = Path("./converted_files")
PROCESSED_DIR.mkdir(exist_ok=True)

# 📚 Knowledge Base
KB_DIR = Path("./knowledge_base")
KB_DIR.mkdir(exist_ok=True)

# 💾 Index file for tracking processed documents
INDEX_FILE = KB_DIR / "document_index.json"

def load_document_index():
    """Load the document index or create it if it doesn't exist."""
    if not INDEX_FILE.exists():
        return {"documents": []}
    
    try:
        with open(INDEX_FILE, 'r') as f:
            return json.load(f)
    except (json.JSONDecodeError, FileNotFoundError):
        print(f"⚠️ Error loading index file. Creating a new one.")
        return {"documents": []}

def save_document_index(index_data):
    """Save the document index to disk."""
    with open(INDEX_FILE, 'w') as f:
        json.dump(index_data, f, indent=4)
    print(f"✅ Document index updated with {len(index_data['documents'])} documents")

def register_document(pdf_path, txt_path, content, chunks=None):
    """Register a processed document in the knowledge base index."""
    # Load existing index
    index_data = load_document_index()
    
    # Create document record
    document_id = f"doc_{len(index_data['documents']) + 1}"
    
    # Generate embeddings and chunks for search
    chunks = chunks or chunk_text(content)
    chunk_embeddings = []
    
    for i, chunk in enumerate(chunks):
        try:
            # Generate embedding for chunk - using 'input' instead of 'prompt'
            embedding_result = ollama.embed(model=EMBEDDING_MODEL, input=chunk)
            embedding = embedding_result.get('embedding', [])
            
            # Store chunk with its embedding
            chunk_file = KB_DIR / f"{document_id}_chunk_{i}.json"
            chunk_data = {
                "text": chunk,
                "embedding": embedding
            }
            
            with open(chunk_file, 'w') as f:
                json.dump(chunk_data, f)
            
            # Remember the chunk reference
            chunk_embeddings.append({
                "chunk_id": f"{document_id}_chunk_{i}",
                "file": str(chunk_file)
            })
        except Exception as e:
            print(f"⚠️ Error embedding chunk {i}: {e}")
    
    # Add document to index
    doc_entry = {
        "id": document_id,
        "original": str(pdf_path),
        "text_path": str(txt_path),
        "processed_date": datetime.now().isoformat(),
        "chunks": chunk_embeddings
    }
    
    index_data["documents"].append(doc_entry)
    save_document_index(index_data)
    return document_id

def chunk_text(text, chunk_size=512, overlap=100):
    """Split text into overlapping chunks for better semantic search."""
    # Split into paragraphs first
    paragraphs = [p for p in text.split('\n\n') if p.strip()]
    
    chunks = []
    current_chunk = []
    current_length = 0
    
    for para in paragraphs:
        # If adding this paragraph would exceed chunk size, finalize current chunk
        if current_length + len(para) > chunk_size and current_chunk:
            chunks.append('\n\n'.join(current_chunk))
            # Keep some overlap by retaining the last paragraph if possible
            if len(current_chunk) > 1:
                current_chunk = current_chunk[-1:]
                current_length = len(current_chunk[0])
            else:
                current_chunk = []
                current_length = 0
        
        # Add current paragraph to the chunk
        current_chunk.append(para)
        current_length += len(para)
    
    # Add the last chunk if not empty
    if current_chunk:
        chunks.append('\n\n'.join(current_chunk))
    
    return chunks

def list_documents():
    """List all processed documents in the knowledge base."""
    index_data = load_document_index()
    if not index_data["documents"]:
        print("📚 No documents have been processed yet.")
        return
    
    print(f"\n📚 Knowledge Base: {len(index_data['documents'])} documents")
    print("-" * 60)
    for i, doc in enumerate(index_data["documents"]):
        print(f"{i+1}. {Path(doc['original']).name}")
        print(f"   ID: {doc['id']} | Processed: {doc['processed_date'][:10]}")
        print(f"   Chunks: {len(doc['chunks'])}")
        print("-" * 60)

def get_similar_chunks(query, top_k=3):
    """Find text chunks similar to the query using vector similarity."""
    try:
        # Get embedding for the query - using 'input' instead of 'prompt'
        query_result = ollama.embed(model=EMBEDDING_MODEL, input=query)
        query_embedding = query_result.get('embedding', [])
        
        if not query_embedding:
            print("⚠️ Failed to generate query embedding")
            return []
        
        # Load document index
        index_data = load_document_index()
        
        # Check if we have documents
        if not index_data["documents"]:
            print("📚 No documents in knowledge base yet")
            return []
        
        # Collect all chunks with their embeddings
        chunks_with_similarity = []
        
        for doc in index_data["documents"]:
            for chunk_ref in doc["chunks"]:
                try:
                    # Load chunk data
                    chunk_file = Path(chunk_ref["file"])
                    if chunk_file.exists():
                        with open(chunk_file, 'r') as f:
                            chunk_data = json.load(f)
                            
                        # Calculate cosine similarity
                        chunk_embedding = chunk_data.get("embedding", [])
                        if chunk_embedding:
                            similarity = cosine_similarity(query_embedding, chunk_embedding)
                            
                            chunks_with_similarity.append({
                                "doc_id": doc["id"],
                                "doc_name": Path(doc["original"]).name,
                                "text": chunk_data["text"],
                                "similarity": similarity
                            })
                except Exception as e:
                    print(f"⚠️ Error processing chunk {chunk_ref['chunk_id']}: {e}")
        
        # Sort by similarity (highest first) and get top k
        chunks_with_similarity.sort(key=lambda x: x["similarity"], reverse=True)
        return chunks_with_similarity[:top_k]
    
    except Exception as e:
        print(f"⚠️ Error searching knowledge base: {e}")
        return []

def cosine_similarity(vec1, vec2):
    """Calculate cosine similarity between two vectors."""
    if not vec1 or not vec2:
        return 0
    
    vec1 = np.array(vec1)
    vec2 = np.array(vec2)
    
    dot_product = np.dot(vec1, vec2)
    norm_a = np.linalg.norm(vec1)
    norm_b = np.linalg.norm(vec2)
    
    if norm_a == A or norm_b == 0:
        return 0
    
    return dot_product / (norm_a * norm_b)

def embed_text(text):
    """Embeds text using Ollama without storing vector data in files."""
    try:
        # Using 'input' instead of 'prompt'
        response = ollama.embed(model=EMBEDDING_MODEL, input=text)
        return text  # Return the original text for saving to files
    except Exception as e:
        print(f"❌ Embedding error: {e}")
        return None

def extract_text_from_pdf(pdf_path):
    """Extracts text from a PDF file using multiple methods for better reliability."""
    pdf_path_str = str(pdf_path)
    extracted_text = ""
    
    # Method 1: Try PyPDF2 first
    try:
        print("🔍 Attempting extraction with PyPDF2...")
        reader = PdfReader(pdf_path_str)
        pypdf_text = ""
        
        for page_num, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text:
                pypdf_text += f"{page_text}\n\n"
            else:
                print(f"⚠️ PyPDF2: No text extracted from page {page_num + 1}")
        
        if pypdf_text.strip():
            print(f"✅ PyPDF2 successfully extracted {len(pypdf_text)} characters")
            extracted_text = pypdf_text
        else:
            print("⚠️ PyPDF2 extraction yielded no text, trying alternative method...")
    except Exception as e:
        print(f"⚠️ PyPDF2 extraction error: {e}")
    
    # Method 2: If PyPDF2 failed or returned no text, try pdftotext if available
    if not extracted_text:
        try:
            print("🔍 Attempting extraction with pdftotext (if installed)...")
            with tempfile.NamedTemporaryFile(suffix='.txt') as temp:
                # Try to use pdftotext (from poppler-utils) if installed
                result = subprocess.run(
                    ['pdftotext', '-layout', pdf_path_str, temp.name],
                    capture_output=True,
                    text=True
                )
                
                if result.returncode == 0:
                    with open(temp.name, 'r', encoding='utf-8') as f:
                        pdftotext_text = f.read()
                    
                    if pdftotext_text.strip():
                        print(f"✅ pdftotext successfully extracted {len(pdftotext_text)} characters")
                        extracted_text = pdftotext_text
                    else:
                        print("⚠️ pdftotext extraction yielded no text")
                else:
                    print(f"⚠️ pdftotext error: {result.stderr}")
        except FileNotFoundError:
            print("⚠️ pdftotext not found on system, skipping alternative extraction")
        except Exception as e:
            print(f"⚠️ Alternative extraction error: {e}")
    
    # Check if we have any text after trying all methods
    if not extracted_text:
        print("❌ Failed to extract text with all available methods")
        return ""
    
    # Process the text to make it more readable
    processed_text = ""
    pages = extracted_text.split("\f")  # Form feed character often separates PDF pages
    
    for page_num, page_content in enumerate(pages):
        if page_content.strip():
            processed_text += f"--- Page {page_num + 1} ---\n\n{page_content.strip()}\n\n"
    
    return processed_text.strip()

def process_pdf(pdf_path):
    """Extracts text from PDF, embeds it, and saves clean conversions."""
    pdf_path = Path(pdf_path).resolve()
    if not pdf_path.exists():
        print(f"❌ Error: File not found → {pdf_path}")
        return

    print(f"📄 Processing '{pdf_path.name}'...")

    # Extract text from PDF using multiple methods
    extracted_text = extract_text_from_pdf(pdf_path)

    if not extracted_text:
        print(f"❌ Failed to extract text from {pdf_path.name}")
        print("💡 This PDF might be:")
        print("   - Scanned (image-based) without OCR")
        print("   - Protected/encrypted")
        print("   - Using non-standard fonts")
        print("   - Corrupted or malformed")
        return

    # Debug: Show a sample of the extracted text
    sample_length = min(200, len(extracted_text))
    print(f"📊 Extracted {len(extracted_text)} characters from {pdf_path.name}")
    print(f"📃 Sample of extracted text: \n{extracted_text[:sample_length]}...")

    # Prepare a clean version of text without the page markers for better readability
    clean_text = extracted_text
    if "--- Page" in clean_text:
        # Remove page markers if present but preserve content
        clean_text = "\n\n".join([
            line for line in clean_text.split("\n")
            if not line.strip().startswith("--- Page")
        ])

    # Get the PDF filename without extension for use in the document title and headings
    pdf_filename = pdf_path.stem
    
    # Save TXT - preserving the original filename
    txt_path = PROCESSED_DIR / f"{pdf_filename}.txt"
    with open(txt_path, "w", encoding="utf-8") as txt_file:
        # Add the PDF filename as the first line of the text file
        txt_file.write(f"# {pdf_filename}\n\n")
        txt_file.write(clean_text)
    print(f"✅ Saved TXT → {txt_path}")

    # Save Markdown - preserving the original filename
    md_path = PROCESSED_DIR / f"{pdf_filename}.md"
    with open(md_path, "w", encoding="utf-8") as md_file:
        # Creating proper markdown with the PDF filename as the title
        md_content = f"# {pdf_filename}\n\n{clean_text}"
        md_file.write(md_content)
    print(f"✅ Saved Markdown → {md_path}")

    # Save DOCX - preserving the original filename
    docx_path = PROCESSED_DIR / f"{pdf_filename}.docx"
    doc = docx.Document()
    
    # Add the PDF filename as the document title/heading
    doc.add_heading(pdf_filename, level=1)
    
    # Split text into paragraphs for better DOCX formatting
    paragraphs = clean_text.split('\n\n')
    for para in paragraphs:
        if para.strip():
            doc.add_paragraph(para.strip())
    
    doc.save(docx_path)
    print(f"✅ Saved DOCX → {docx_path}")
    
    # Add to knowledge base for chat capability
    print("🧠 Adding document to knowledge base...")
    chunks = chunk_text(clean_text)
    print(f"📊 Document divided into {len(chunks)} semantic chunks")
    
    doc_id = register_document(pdf_path, txt_path, clean_text, chunks)
    print(f"✅ Document registered with ID: {doc_id}")

    print(f"🎯 Completed processing {pdf_path.name}")

def process_directory(dir_path):
    """Processes all PDFs in a given directory."""
    dir_path = Path(dir_path).expanduser().resolve()
    
    print(f"🔍 Debug: Checking path → {dir_path}")

    if not dir_path.exists() or not dir_path.is_dir():
        print(f"❌ Error: Directory not found → {dir_path}")
        return

    pdf_files = list(dir_path.glob("*.pdf"))
    if not pdf_files:
        print(f"⚠️ No PDFs found in {dir_path}")
        return

    print(f"📂 Found {len(pdf_files)} PDFs. Processing...")

    for pdf in pdf_files:
        process_pdf(pdf)

    print("✅ All PDFs processed!")

def chat():
    """Starts an interactive chat with embedded documents."""
    index_data = load_document_index()
    if not index_data["documents"]:
        print("📚 No documents in knowledge base yet. Process a PDF first.")
        return
    
    print("\n💬 Chat with your Documents")
    print("Type 'exit' to return to main menu")
    print(f"Knowledge base contains {len(index_data['documents'])} documents")
    
    chat_history = []
    
    while True:
        user_query = input("\n🙋 You: ").strip()
        
        if user_query.lower() == 'exit':
            print("Returning to main menu...")
            break
        
        if not user_query:
            continue
        
        # Find relevant document chunks
        print("🔍 Searching knowledge base...")
        relevant_chunks = get_similar_chunks(user_query)
        
        if not relevant_chunks:
            print("❌ No relevant information found in the knowledge base.")
            continue
        
        # Format context from relevant chunks
        context = "\n\n".join([
            f"From document '{chunk['doc_name']}':\n{chunk['text']}" 
            for chunk in relevant_chunks
        ])
        
        # Build prompt with context and chat history
        system_prompt = "You are Local Doc AI, a helpful assistant that answers questions based on the user's documents. Only use information from the provided document context. If you don't know or the answer isn't in the context, say so."
        
        history_text = ""
        if chat_history:
            history_text = "\n\n".join([
                f"Previous Q: {q}\nPrevious A: {a}" 
                for q, a in chat_history[-3:]  # Include last 3 exchanges for context
            ])
            history_text = f"\nCHAT HISTORY:\n{history_text}\n"
        
        prompt = f"{system_prompt}\n\nCONTEXT FROM DOCUMENTS:\n{context}\n{history_text}\nQUESTION: {user_query}\nANSWER:"
        
        # Generate response using LLM
        print("🤖 Generating response...")
        try:
            response = ollama.chat(
                model=CHAT_MODEL,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": f"CONTEXT: {context}\n\nQUESTION: {user_query}"}
                ]
            )
            
            answer = response['message']['content']
            
            # Display response
            print(f"\n🤖 AI: {answer}")
            
            # Update chat history
            chat_history.append((user_query, answer))
            
            # Show source documents
            print("\n📚 Sources:")
            for i, chunk in enumerate(relevant_chunks[:3]):
                print(f"  {i+1}. {chunk['doc_name']} (relevance: {chunk['similarity']:.2f})")
            
        except Exception as e:
            print(f"❌ Error generating response: {e}")

def check_dependencies():
    """Checks for the presence of external tools that might help with PDF processing."""
    print("🔍 Checking for helpful dependencies...")
    
    # Check for pdftotext (from Poppler utils)
    try:
        result = subprocess.run(['pdftotext', '-v'], 
                               stdout=subprocess.PIPE, 
                               stderr=subprocess.PIPE, 
                               text=True)
        if result.returncode == 0:
            version_info = result.stderr.strip() if result.stderr else result.stdout.strip()
            print(f"✅ pdftotext found: {version_info}")
        else:
            print("❌ pdftotext is not working properly")
    except FileNotFoundError:
        print("❌ pdftotext not found. For better PDF extraction, consider installing:")
        print("   - Linux: sudo apt-get install poppler-utils")
        print("   - macOS: brew install poppler")
        print("   - Windows: Install from http://blog.alivate.com.au/poppler-windows/")
    
    # Check PyPDF2 version
    import pkg_resources
    try:
        pypdf_version = pkg_resources.get_distribution("PyPDF2").version
        print(f"✅ PyPDF2 version: {pypdf_version}")
    except pkg_resources.DistributionNotFound:
        print("❌ PyPDF2 not found in installed packages")
    
    # Check for OCRmyPDF for potential enhancement
    try:
        result = subprocess.run(['ocrmypdf', '--version'], 
                               stdout=subprocess.PIPE, 
                               stderr=subprocess.PIPE, 
                               text=True)
        if result.returncode == 0:
            version_info = result.stdout.strip()
            print(f"✅ OCRmyPDF found: {version_info}")
            print("   This can be used to add OCR to scanned PDFs if needed")
        else:
            print("❌ OCRmyPDF is not working properly")
    except FileNotFoundError:
        print("ℹ️ OCRmyPDF not found (optional for OCR capability)")
    
    # Check Ollama availability
    try:
        result = ollama.list()
        # Display available models
        models = result.get('models', [])
        model_names = [model.get('name', 'unknown') for model in models]
        
        if EMBEDDING_MODEL in model_names:
            print(f"✅ Embedding model '{EMBEDDING_MODEL}' is available")
        else:
            print(f"⚠️ Embedding model '{EMBEDDING_MODEL}' not found in Ollama")
            print(f"   Run 'ollama pull {EMBEDDING_MODEL}' to download it")
        
        if CHAT_MODEL in model_names:
            print(f"✅ Chat model '{CHAT_MODEL}' is available")
        else:
            print(f"⚠️ Chat model '{CHAT_MODEL}' not found in Ollama")
            print(f"   Run 'ollama pull {CHAT_MODEL}' to download it")
            
    except Exception as e:
        print(f"❌ Ollama not available or error connecting: {e}")
        print("   Make sure Ollama is installed and running")
        
    print("\n💡 Missing tools can be installed to improve PDF processing capabilities")

def main():
    """Command-line interface."""
    print("🚀 Welcome to Local Doc AI")
    print(COMMANDS)
    
    # Quick dependency check on startup
    print("\nℹ️ Run 'check_deps' for detailed dependency information")

    while True:
        command = input("\n⚡ Enter command: ").strip().split()

        if not command:
            continue

        action = command[0]
        arg = " ".join(command[1:]) if len(command) > 1 else None

        if action == "open_pdf" and arg:
            process_pdf(arg)
        elif action == "open_dir" and arg:
            process_directory(arg)
        elif action == "chat":
            chat()
        elif action == "list_docs":
            list_documents()
        elif action == "check_deps":
            check_dependencies()
        elif action == "exit":
            print("👋 Exiting. See you next time!")
            break
        else:
            print("⚠️ Invalid command. Try again.")
            print(COMMANDS)

if __name__ == "__main__":
    main()
